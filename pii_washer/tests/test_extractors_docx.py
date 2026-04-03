import io

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from pii_washer.extractors.docx import DocxExtractor


# === Helpers ===


def _make_docx(**kwargs) -> bytes:
    """Build a .docx file in memory and return raw bytes.

    Accepts keyword arguments that control what content is added:
        paragraphs: list[str]       — regular body paragraphs
        headings:   list[tuple]     — (level: int, text: str)
        bullets:    list[str]       — bullet list items
        numbers:    list[str]       — numbered list items
        tables:     list[list[list[str]]] — list of tables, each a list of rows,
                                            each row a list of cell strings
    """
    doc = Document()

    for text in kwargs.get("paragraphs", []):
        doc.add_paragraph(text)

    for level, text in kwargs.get("headings", []):
        doc.add_heading(text, level=level)

    for text in kwargs.get("bullets", []):
        doc.add_paragraph(text, style="List Bullet")

    for text in kwargs.get("numbers", []):
        doc.add_paragraph(text, style="List Number")

    for rows in kwargs.get("tables", []):
        if not rows:
            continue
        col_count = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=col_count)
        for ri, row in enumerate(rows):
            for ci, cell_text in enumerate(row):
                table.cell(ri, ci).text = cell_text

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_raw(doc: Document) -> bytes:
    """Serialise a pre-built Document object to bytes."""
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def extractor():
    return DocxExtractor()


# === Happy path — content types ===


def test_returns_string_type(extractor):
    content = _make_docx(paragraphs=["Hello world"])
    result = extractor.extract(content, "test.docx")
    assert isinstance(result, str)


def test_paragraph_extraction(extractor):
    content = _make_docx(paragraphs=["First paragraph.", "Second paragraph."])
    result = extractor.extract(content, "test.docx")
    assert "First paragraph." in result
    assert "Second paragraph." in result


def test_paragraph_separation(extractor):
    content = _make_docx(paragraphs=["Alpha", "Beta"])
    result = extractor.extract(content, "test.docx")
    # Paragraphs must be separated by a blank line (double newline)
    assert "Alpha\n\nBeta" in result


def test_heading_extraction(extractor):
    content = _make_docx(headings=[(1, "Section One"), (2, "Subsection")])
    result = extractor.extract(content, "test.docx")
    assert "Section One" in result
    assert "Subsection" in result


def test_bullet_list_extraction(extractor):
    content = _make_docx(bullets=["First item", "Second item", "Third item"])
    result = extractor.extract(content, "test.docx")
    assert "- First item" in result
    assert "- Second item" in result
    assert "- Third item" in result


def test_numbered_list_extraction(extractor):
    content = _make_docx(numbers=["Step one", "Step two", "Step three"])
    result = extractor.extract(content, "test.docx")
    assert "1. Step one" in result
    assert "2. Step two" in result
    assert "3. Step three" in result


def test_table_extraction(extractor):
    rows = [["Name", "Email"], ["Alice", "alice@example.com"]]
    content = _make_docx(tables=[rows])
    result = extractor.extract(content, "test.docx")
    assert "| Name | Email |" in result
    assert "| Alice | alice@example.com |" in result


def test_table_pipe_delimited_format(extractor):
    """Each row starts and ends with a pipe character."""
    rows = [["A", "B", "C"]]
    content = _make_docx(tables=[rows])
    result = extractor.extract(content, "test.docx")
    assert "| A | B | C |" in result


def test_mixed_content_order_preserved(extractor):
    """Headings, paragraphs, lists, and tables all appear in document order."""
    doc = Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Intro paragraph.")
    doc.add_paragraph("Bullet one", style="List Bullet")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Col1"
    table.cell(0, 1).text = "Col2"
    content = _make_docx_raw(doc)

    result = extractor.extract(content, "mixed.docx")
    title_pos = result.index("Title")
    intro_pos = result.index("Intro paragraph.")
    bullet_pos = result.index("- Bullet one")
    table_pos = result.index("| Col1 | Col2 |")
    assert title_pos < intro_pos < bullet_pos < table_pos


# === Error handling ===


def test_empty_document_raises_value_error(extractor):
    doc = Document()
    content = _make_docx_raw(doc)
    with pytest.raises(ValueError, match="No text content"):
        extractor.extract(content, "empty.docx")


def test_corrupted_file_raises_value_error(extractor):
    garbage = b"This is not a valid docx file at all \x00\x01\x02"
    with pytest.raises(ValueError, match="could not be read"):
        extractor.extract(garbage, "bad.docx")


def test_empty_bytes_raises_value_error(extractor):
    with pytest.raises(ValueError, match="could not be read"):
        extractor.extract(b"", "empty.docx")
